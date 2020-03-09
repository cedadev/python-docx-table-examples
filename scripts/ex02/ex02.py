import collections
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dreqPy import dreq
from dreqPy.extensions import collect
dq = dreq.loadDreq()
collect.add( dq )

print ( 'VERSION: %s' % dq.version )

class DocxTable(object):
  def __init__(self,ofn):
    self.doc = docx.Document('Doc1.docx')
    self.doc.save( ofn )
    self.ofn = ofn
    self.pos = 2
    self.nrow = 3
    self.table = self.doc.tables[0]

  def addRow(self,ll,notes=False):
    """Add a row of test values to a table, provided as a list.

    notes: if set True, add an additional row, merging cells after the first.
    """

    self.pos += 1
    if self.pos > self.nrow:
      self.table.add_row()
      self.nrow += 1
    for k in range( len(ll) ):
      self.table.cell(self.pos-1, k).text = str(ll[k])

    if notes:
      self.table.add_row()
      self.nrow += 1
      self.pos += 1
      c = self.table.cell(self.pos-1, 0)
      p = c.add_paragraph( 'Notes' )
      p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
      a = self.table.cell(self.pos-1,1)
      b = self.table.cell(self.pos-1,len(ll)-1)
      a.merge(b)
    self.doc.save( self.ofn )

  def close(self):
    self.doc.save( self.ofn )
      


class Stabs(object):
  defaults = { 'uid':'Identifier, unique within a given version of the data request.',
               'label':'A single word,  with restricted character set. Specialization of SKOS prefLabel.',
               'title':'A few words describing the object. Specialization of Dublin Core title.',
               'description':'An extended description of the object/concept. Specialization of SKOS definition.' }
  def __init__(self,dq):
    self.dq = dq
    self.es = dict()

## load in description ... which doesn't make it into the main structure ... yet
    self.es2 = dict()
    for t in dq.c.ttl2:
      self.es2[t.uid] = dq.c.__desc__.get( t.label, '' )

    self.cc = collections.defaultdict(dict)
    for i in dq.coll['__sect__'].items:
      self.es[i.label] = i
    for i in dq.coll['__main__'].items:
      a,b = i.uid.split( '::' )
      x,y = b.split('.')
      self.cc[x][y] = i

    ee = dict()
    oo = open( 'inputs.texlet', 'w' )
    for s in dq.coll.keys():
      if s[0] != '_':
        if s in self.cc:
          ee[ self.es[s].title ] = s
        else:
          print ('SECTION NOT FOUND: %s' % s )

    ns = len( ee.keys() )
    na = sum( [len(self.cc[sect].keys()) for sect in self.cc] )
    for t in sorted( ee.keys() ):
          s = ee[t]
          oo.write( '\\input{slist_%s.texlet}\n' % s )
          self.psect(s)
    oo.close()
    print ( 'ns=%s, na=%s' % (ns,na) )

  def _latex(self,x):
    return x.replace( '_', '\\_' )

  def psect(self,sect):
    assert sect in self.es, 'Section not found in __sect__: %s' % sect
    assert sect in self.cc, 'Section not found in __main__: %s' % sect
    print (self.es[sect].title)
    print (self.cc[sect].keys())
    oo = open( 'slist_%s.texlet' % sect, 'w' )
    oo.write( '\\section*{%s [%s]}\n' % (self.es[sect].title,self.es[sect].label) )
    
    oo.write( '%s\n\n' % self._latex(self.dq.c._desc[self.es[sect].label] ) )

    f = lambda x: '%s: %s' % ({'label':'00','title':'01','description':'02','uid':'03'}.get(x,'99'),x)

    oo.write( '\\sectionTable{\n' )
    ox = DocxTable( 'docxDocs/%s.docx' % sect )
    for k in sorted( self.cc[sect].keys(), key=f ):
      this = self.cc[sect][k]

      dd = this.description.strip()
      print ( '**** %s, %s:: [%s]' % (sect,this.label,this.description) )
      thatx = [this.label, this.title, this.description, this.type]
      if this.label in self.defaults and dd == '':
        this.description = '{\\it %s}' % self.defaults[k]
        thatx[2] = '%s [DEF]' % self.defaults[k]
        print ( '--> %s' % this.description )

      lab = this.label
      desc = thatx[2]
      if this.useClass == None:
        clss = ''
      else:
        clss = this.useClass.strip()

      lab += ' [%s]' % thatx[3]
      if clss != '':
        lab += ' {%s}' % clss
      if not this.required:
        lab += ' (OPT)'
      thatx[0] = lab
      ##if usage != '':
        ##desc += ' [%s]' % usage
        ##thatx[2] = desc

      that = [this.label, this.title, this.description, '\\ttt{%s}' % this.type, '\\ttt{%s}' % this.uid ]
      line = '\\hline\n' + ' & '.join( [self._latex(x) for x in that[:4]] ) + ' \\\\ \n'
      oo.write( line )
      withNotes = sect in ['grids']
      ox.addRow( thatx[:3], notes=withNotes )
       
    oo.write( '}% End of sectionTable argument\n' )
    oo.close()
    ox.close()

    
    

for k in sorted( dq.coll.keys() ):
  if len( dq.coll[k].items ) > 0:
    i = dq.coll[k].items[0]
    print ( '%s & %s & %s &  \\' % (k,len( dq.coll[k].items ), i._h.title ) )


nn = 0
for i in dq.coll['requestLink'].items:
  n1 = len( dq.inx.iref_by_sect[i.refid].a['requestVar'] )
  n2 = len( i._get__expt() )
  n3 = len( dq.inx.iref_by_sect[i.uid].a['objectiveLink'] )
  nn += n1*n2*n3

nnn = 1
for y in [len(dq.coll[x].items) for x in ['CMORvar','experiment','objective']]:
  nnn = nnn*y

print ( 'Virtual requests (var-expt-objective): %s [%s, %4.2f%%]' % (nn, nnn, float(100*nn)/nnn) )

nn = 0
ee = dict()
for i in dq.coll['experiment'].items:
  n1 = len( i._get__CMORvar() )
  ee[i.label] = n1
  nn += n1

ks =  sorted( ee.keys(), key=lambda x: ee[x], reverse=True)


nnn = 1
for y in [len(dq.coll[x].items) for x in ['CMORvar','experiment']]:
  nnn = nnn*y

print ( 'Virtual requests (var-expt): %s [%s, %4.2f%%]' % (nn, nnn, float(100*nn)/nnn) )

c1 = [x for x in dq.coll['CMORvar'].items  if 'requestVar' not in dq.inx.iref_by_sect[x].a]
nnn = len( dq.coll['CMORvar'].items )
print ( 'CMORVvars: %s [%s]' % (nnn,len(c1)) )
for k in ks[:5]:
  print ( 'Var count (top): %s [%s, %4.2f%%]' % (k,ee[k], float(100*ee[k])/nnn) )
for k in ks[-5:]:
  print ( 'Var count (bottom): %s [%s, %4.2f%%]' % (k,ee[k], float(100*ee[k])/nnn) )

ss = Stabs(dq)
